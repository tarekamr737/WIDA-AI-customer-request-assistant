from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter

from src.file_parser import FileParseError, extract_request_text


def test_extracts_utf8_txt_upload() -> None:
    content = "طلب عربي من عميل".encode("utf-8")

    assert extract_request_text("request.TXT", content) == "طلب عربي من عميل"


def test_extracts_docx_paragraphs() -> None:
    document = Document()
    document.add_paragraph("السطر الأول")
    document.add_paragraph("السطر الثاني")
    buffer = BytesIO()
    document.save(buffer)

    assert extract_request_text("request.docx", buffer.getvalue()) == (
        "السطر الأول\nالسطر الثاني"
    )


def test_image_only_pdf_has_an_ocr_specific_error() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    buffer = BytesIO()
    writer.write(buffer)

    with pytest.raises(FileParseError, match="OCR"):
        extract_request_text("scan.pdf", buffer.getvalue())


def test_empty_and_unsupported_uploads_are_rejected() -> None:
    with pytest.raises(FileParseError, match="فارغ"):
        extract_request_text("request.txt", b"")
    with pytest.raises(FileParseError, match="غير مدعوم"):
        extract_request_text("request.csv", b"value")
